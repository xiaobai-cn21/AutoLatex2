"""生成用于 docx 解析测试的示例文件。"""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document  # type: ignore
from docx.enum.style import WD_STYLE_TYPE  # type: ignore
from docx.shared import Inches  # type: ignore
from PIL import Image, ImageDraw  # type: ignore


BASE_DIR = Path(__file__).resolve().parent / "docx_samples"
ASSETS_DIR = BASE_DIR / "assets"


def _ensure_dirs() -> None:
    BASE_DIR.mkdir(parents=True, exist_ok=True)
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)


def _create_sample_image() -> Path:
    image_path = ASSETS_DIR / "sample_figure.png"
    if image_path.exists():
        return image_path
    img = Image.new("RGB", (600, 300), color=(46, 134, 222))
    draw = ImageDraw.Draw(img)
    draw.rectangle((50, 80, 550, 220), outline="white", width=5)
    draw.text((70, 130), "AutoTeX Diagram", fill="white")
    img.save(image_path)
    return image_path


def _ensure_code_style(document: Document) -> None:
    styles = document.styles
    if "Code" in styles:
        return
    style = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = "Consolas"
    font.size = document.styles["Normal"].font.size


def create_sample_paper_full() -> None:
    doc = Document()
    _ensure_code_style(doc)

    doc.add_paragraph("AutoTeX Research Paper").style = "Title"
    doc.add_paragraph(
        "Alice Zhang (Tsinghua University); Bob Li (MIT); Carol Chen (HKUST)"
    )
    doc.add_paragraph("Abstract: This document demonstrates the AutoTeX parsing flow.")
    doc.add_paragraph("Keywords: AutoTeX, Parsing, Schema, Word, Pipeline")

    doc.add_heading("1 Introduction", level=1)
    doc.add_paragraph(
        "This paper introduces the AutoTeX pipeline that converts complex documents "
        "into structured JSON outputs for downstream agents. The system leverages "
        "AI techniques and strict schema validation [1]."
    )

    doc.add_heading("1.1 Background", level=2)
    doc.add_paragraph(
        "Traditional document digitization workflows rely on manual formatting. "
        "AutoTeX automates this process by combining deterministic parsing with "
        "LLM-based inference."
    )

    doc.add_paragraph("$$E = mc^2$$")

    code_para = doc.add_paragraph()
    code_para.style = "Code"
    code_para.add_run("for i in range(3):\n    print('AutoTeX rocks!')")

    doc.add_paragraph("Table 1: System Metrics")
    table = doc.add_table(rows=4, cols=3)
    table.style = "Light Grid Accent 1"
    headers = ["Module", "Latency (ms)", "Accuracy"]
    for idx, text in enumerate(headers):
        table.cell(0, idx).text = text
    data_rows = [
        ["Parser", "32", "0.91"],
        ["OCR", "58", "0.87"],
        ["Renderer", "20", "0.95"],
    ]
    for row_idx, row_data in enumerate(data_rows, start=1):
        for col_idx, value in enumerate(row_data):
            table.cell(row_idx, col_idx).text = value

    doc.add_paragraph("Figure 1: AutoTeX Architecture")
    image_path = _create_sample_image()
    doc.add_picture(str(image_path), width=Inches(4))

    doc.add_paragraph("Further explanation with another paragraph referencing [2].")

    doc.add_heading("References", level=1)
    doc.add_paragraph(
        "[1] Alice Zhang, Bob Li. AutoTeX System Overview. Journal of Automation, 2024."
    )
    doc.add_paragraph(
        "[2] Carol Chen. Structured Document Understanding. Conference on AI Systems, 2025."
    )

    doc.save(BASE_DIR / "sample_paper_full.docx")


def create_sample_paper_min() -> None:
    doc = Document()
    doc.add_paragraph("Minimal AutoTeX Paper").style = "Title"
    doc.add_paragraph("Dana Wu (OpenAI); Eric Sun (DeepMind)")
    doc.add_paragraph("Abstract: Minimal example to verify mandatory fields.")
    doc.add_paragraph(
        "This minimal document ensures that the parser can extract required metadata and one paragraph."
    )
    doc.save(BASE_DIR / "sample_paper_min.docx")


def create_sample_paper_no_bib() -> None:
    doc = Document()
    doc.add_paragraph("AutoTeX Without Bibliography").style = "Title"
    doc.add_paragraph("Fiona Zhang (CMU)")
    doc.add_paragraph("Abstract: Document without bibliography for edge case testing.")
    doc.add_heading("1 Discussion", level=1)
    doc.add_paragraph(
        "Edge cases are essential to ensure that AutoTeX handles missing sections gracefully."
    )
    doc.save(BASE_DIR / "sample_paper_no_bib.docx")


def main() -> None:
    _ensure_dirs()
    _create_sample_image()
    create_sample_paper_full()
    create_sample_paper_min()
    create_sample_paper_no_bib()
    print(f"Sample DOCX files created under {BASE_DIR}")


if __name__ == "__main__":
    main()


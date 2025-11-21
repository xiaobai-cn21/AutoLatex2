# filename: src/autolatex/data_synthesis/paper_generator.py

import os
import re
from tqdm import tqdm
import docx
from openai import OpenAI

# --- 1. 配置区 (保持不变) ---
API_KEY = "sk-8f4a9f0320c645238be600f6b7911fbf"

client = OpenAI(
    api_key=API_KEY,
    base_url="https://api.deepseek.com/v1"
)

# --- 2. Prompt工程 (保持不变) ---
RESEARCHER_PROMPT_TEMPLATE = """
你是一名顶尖的{field}领域的研究员，你的任务是撰写一篇关于“{topic}”的高质量伪论文。

请严格遵守以下规则：
1.  **使用Markdown语法进行排版**：
    - 论文标题使用一级标题 `# 标题`。
    - 摘要前写 `**摘要**：`。
    - 章节标题使用二级标题 `## 1. 引言`。
    - 无序列表使用 `-` 或 `*`。
2.  **内容要求**：论文必须包含标题、摘要、至少三个章节（例如：引言、方法论、实验、结论）。内容要显得专业、连贯。
3.  **【最重要】占位符规则**：当你需要展示一个复杂的公式或一个表格时，**绝对不要**自己编写LaTeX代码或Markdown表格。请在正文中插入一个清晰的、描述性的占位符：
    - **公式占位符格式**：`[FORMULA: 这里是公式的自然语言描述]`
    - **表格占位符格式**：`[TABLE: 这里是表格的自然语言描述]`

**示例**：
...我们的核心方法基于二次方程求解，其公式为 [FORMULA: x等于2a分之负b加减根号下b的平方减4ac]。我们将不同模型的结果总结在下表中 [TABLE: 一个3行4列表格，列为模型名称、准确率、召回率和F1分数，行是BERT、RoBERTA和OurModel的具体数值]。

现在，请开始撰写你的论文。
"""


# --- generate_paper_markdown 和 save_paper_to_files 函数 (保持不变) ---

def generate_paper_markdown(field: str, topic: str) -> str:
    """调用LLM生成一篇伪论文的Markdown原文。"""
    prompt = RESEARCHER_PROMPT_TEMPLATE.format(field=field, topic=topic)
    try:
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system",
                 "content": "You are a helpful assistant that writes scientific papers in Markdown format."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2048,
            temperature=0.8,
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"调用API时发生错误: {e}")
        return ""


def save_paper_to_files(paper_content: str, base_filename: str, output_dir: str):
    """将生成的论文内容保存为.md和.docx文件。"""
    os.makedirs(output_dir, exist_ok=True)

    md_path = os.path.join(output_dir, f"{base_filename}.md")
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(paper_content)

    docx_path = os.path.join(output_dir, f"{base_filename}.docx")
    doc = docx.Document()
    for line in paper_content.split('\n'):
        # ... (内部逻辑保持不变)
        line = line.strip()
        if not line: continue
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith(('-', '*', '')):
            item_text = re.sub(r'^[-\*\]\s+', '', line)
            doc.add_paragraph(item_text, style='List Bullet')
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
    doc.save(docx_path)


# ====================== 主要修改在这里 ======================
if __name__ == "__main__":
    research_topics = [
        ("计算机视觉", "基于注意力机制的图像超分辨率重建"),
        ("自然语言处理", "利用知识图谱增强大语言模型的推理能力"),
        ("机器学习理论", "关于深度神经网络泛化能力的理论分析"),
        ("金融科技", "基于区块链的去中心化身份验证系统"),
        ("自动驾驶", "多传感器融合的3D目标检测算法研究"),
    ]

    # <--- 修改点 1：路径现在指向项目根目录下的 data/generated_papers
    # 使用 os.path.abspath 和 .. 来确保路径的健壮性
    # __file__ -> paper_generator.py
    # dirname -> .../data_synthesis/
    # .. -> .../autolatex/
    # .. -> .../src/
    # .. -> 项目根目录
    PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".."))
    output_root_directory = os.path.join(PROJECT_ROOT, "data", "generated_papers")

    print(f"准备生成 {len(research_topics)} 篇伪论文...")
    print(f"文件将保存在: {output_root_directory}")

    for i, (field, topic) in enumerate(tqdm(research_topics, desc="生成论文中")):
        markdown_content = generate_paper_markdown(field, topic)
        if markdown_content:
            # <--- 修改点 2：为每篇论文创建一个独立的子目录 ---
            # 子目录名就是 paper_id，例如 'paper_1', 'paper_2'
            paper_id = f"paper_{i + 1}"
            paper_output_dir = os.path.join(output_root_directory, paper_id)

            # <--- 修改点 3：文件名现在可以简化，因为目录已经唯一了 ---
            # 例如，文件名可以是 'document.md', 'document.docx'
            base_filename = "document"

            # 将 markdown 内容, 文件基名, 和这篇论文专属的输出目录 传给保存函数
            save_paper_to_files(markdown_content, base_filename, paper_output_dir)
        else:
            print(f"未能生成关于 '{topic}' 的论文，已跳过。")

    print("\n✅ 所有伪论文已生成完毕！")
    print(f"文件结构已符合 pipeline.py 的期望。")
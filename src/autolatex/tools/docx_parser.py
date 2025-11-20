"""提供将 Word 文档解析为符合 document_schema 的结构化数据。"""

from __future__ import annotations

import os
import re
import uuid
from typing import Any, Dict, List, Optional

from docx import Document  # type: ignore
from docx.document import Document as DocxDocument  # type: ignore
from docx.oxml.table import CT_Tbl  # type: ignore
from docx.oxml.text.paragraph import CT_P  # type: ignore
from docx.table import _Cell, Table  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore

IMAGES_DIR = os.path.join(os.getcwd(), "parsed_images")


def _ensure_dir(path: str) -> None:
    if not os.path.exists(path):
        os.makedirs(path, exist_ok=True)


def _save_image(image_part) -> str:
    """保存 Word 图片并返回相对路径。"""
    _ensure_dir(IMAGES_DIR)
    image_name = f"{uuid.uuid4()}{image_part.ext}"
    image_path = os.path.join(IMAGES_DIR, image_name)
    with open(image_path, "wb") as image_file:
        image_file.write(image_part.blob)
    return image_path


def _iter_block_items(parent: DocxDocument):
    """遍历文档中的段落和表格（保持原始顺序）。"""
    for child in parent.element.body:
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _guess_heading_level(style_name: str) -> Optional[int]:
    match = re.match(r"Heading\s*(\d)", style_name, re.IGNORECASE)
    if match:
        return int(match.group(1))
    match = re.match(r"标题\s*(\d)", style_name)
    if match:
        return int(match.group(1))
    return None


def _parse_metadata(document: DocxDocument) -> Dict[str, Any]:
    metadata = {"title": "", "authors": [], "abstract": "", "keywords": []}

    # 标题：优先使用 Title 样式；否则第一个非空段落
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style and para.style.name.lower().startswith("title"):
            metadata["title"] = text
            break
        if not metadata["title"]:
            metadata["title"] = text
            break

    author_candidates: List[str] = []
    reached_body = False
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if text.lower().startswith(("abstract", "摘要")):
            reached_body = True
            break
        if text == metadata["title"]:
            continue
        if len(text.split()) <= 20:
            author_candidates.append(text)

    for candidate in author_candidates:
        segments = re.split(r";|、|，|,|\n", candidate)
        for segment in segments:
            segment = segment.strip()
            if len(segment) < 2:
                continue
            name_match = re.match(r"([^(]+)(?:\(([^)]+)\))?", segment)
            if name_match:
                name = name_match.group(1).strip()
                affiliation = name_match.group(2).strip() if name_match.group(2) else ""
                metadata["authors"].append(
                    {"name": name, "affiliation": affiliation, "email": ""}
                )

    abstract_lines: List[str] = []
    keywords: List[str] = []
    collecting_abstract = False
    collecting_keywords = False
    for para in document.paragraphs:
        text = para.text.strip()
        lower_text = text.lower()
        if not text:
            continue
        if lower_text.startswith("abstract") or lower_text.startswith("摘要"):
            collecting_abstract = True
            collecting_keywords = False
            content = text.split(":", 1)[-1].strip() if ":" in text else ""
            if content:
                abstract_lines.append(content)
            continue
        if lower_text.startswith("keywords") or lower_text.startswith("关键词"):
            collecting_keywords = True
            collecting_abstract = False
            keyword_part = text.split(":", 1)[-1]
            keywords.extend([k.strip() for k in re.split(r"[;,，；]", keyword_part) if k.strip()])
            continue
        if collecting_abstract:
            abstract_lines.append(text)
        elif collecting_keywords:
            keywords.extend([k.strip() for k in re.split(r"[;,，；]", text) if k.strip()])

    metadata["abstract"] = "\n".join(abstract_lines).strip()
    metadata["keywords"] = keywords

    if not metadata["title"]:
        raise ValueError("无法从文档中识别标题。")
    if not metadata["authors"]:
        raise ValueError("无法从文档中识别作者信息。")
    if not metadata["abstract"]:
        raise ValueError("无法从文档中识别摘要内容。")

    return metadata


def _paragraph_to_content(para: Paragraph) -> Optional[Dict[str, Any]]:
    text = para.text.strip()
    if not text:
        return None

    if para.style:
        heading_level = _guess_heading_level(para.style.name)
        if heading_level:
            return {"type": "heading", "level": heading_level, "text": text}
        if "code" in para.style.name.lower():
            return {"type": "code", "language": "", "content": text}

    if re.match(r".*\$[^$]+\$.*", text):
        format_type = "inline"
        if text.startswith("$$") and text.endswith("$$"):
            format_type = "display"
        return {"type": "equation", "format": format_type, "latex": text, "image_path": ""}

    reference_matches = re.findall(r"\[[0-9]+\]", text)
    if reference_matches:
        # 返回段落和引用两个内容块
        reference_blocks = [{"type": "reference", "marker": marker} for marker in reference_matches]
        paragraph_block = {"type": "paragraph", "text": text}
        return {"paragraph": paragraph_block, "references": reference_blocks}

    return {"type": "paragraph", "text": text}


def _table_to_content(table: Table) -> Dict[str, Any]:
    table_data: List[List[str]] = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        table_data.append(row_data)
    return {
        "type": "table",
        "caption": "",
        "latex": "",
        "image_path": "",
        "data": table_data,
    }


def parse_docx_to_json(file_path: str) -> Dict[str, Any]:
    document = Document(file_path)
    parsed_data: Dict[str, Any] = {
        "metadata": _parse_metadata(document),
        "content": [],
        "bibliography": [],
    }

    bibliography_started = False

    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            if text.lower().startswith(("references", "参考文献")):
                bibliography_started = True
                continue
            if bibliography_started:
                bib_id = ""
                match = re.match(r"\[(\d+)]", text)
                if match:
                    bib_id = match.group(1)
                parsed_data["bibliography"].append(
                    {
                        "id": bib_id,
                        "type": "misc",
                        "authors": [],
                        "title": "",
                        "venue": "",
                        "year": "",
                        "raw": text,
                    }
                )
                continue

            content_entry = _paragraph_to_content(block)
            if not content_entry:
                continue
            if "paragraph" in content_entry:
                parsed_data["content"].append(content_entry["paragraph"])
                parsed_data["content"].extend(content_entry["references"])
            else:
                parsed_data["content"].append(content_entry)

        elif isinstance(block, Table):
            parsed_data["content"].append(_table_to_content(block))

    if not parsed_data["bibliography"]:
        parsed_data["bibliography"].append(
            {
                "id": "",
                "type": "misc",
                "authors": [],
                "title": "",
                "venue": "",
                "year": "",
                "raw": "",
            }
        )

    return parsed_data


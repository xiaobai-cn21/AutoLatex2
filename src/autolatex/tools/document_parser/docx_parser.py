"""提供将 Word 文档解析为符合 document_schema 的结构化数据。"""

from __future__ import annotations

import os
import re
import uuid
from typing import Any, Dict, List, Optional

from docx import Document  # type: ignore
from docx.document import Document as DocxDocument  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.oxml.table import CT_Tbl  # type: ignore
from docx.oxml.text.paragraph import CT_P  # type: ignore
from docx.table import _Cell, Table  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore

IMAGES_DIR = os.path.join(os.getcwd(), "parsed_images")
CAPTION_PATTERN = re.compile(
    r"^(表|table|图|figure)\s*[\d一二三四五六七八九十0-9\.-]*[:：．.\s-]*(.+)",
    re.IGNORECASE
)


def _ensure_dir(path: str) -> None:
    if not os.path.exists(path):
        os.makedirs(path, exist_ok=True)


def _save_image(image_part) -> str:
    """保存 Word 图片并返回相对路径。"""
    _ensure_dir(IMAGES_DIR)
    
    # 根据 content_type 确定扩展名
    ext_map = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/jpg": ".jpg",
        "image/gif": ".gif",
        "image/bmp": ".bmp",
        "image/tiff": ".tiff",
    }
    ext = ext_map.get(image_part.content_type, ".png")
    
    image_name = f"{uuid.uuid4()}{ext}"
    image_path = os.path.join(IMAGES_DIR, image_name)
    with open(image_path, "wb") as image_file:
        image_file.write(image_part.blob)
    
    # 返回相对路径（相对于项目根目录），统一使用正斜杠以兼容不同操作系统
    relative_path = os.path.join("parsed_images", image_name)
    return relative_path.replace("\\", "/")


def _iter_block_items(parent: DocxDocument):
    """遍历文档中的段落和表格（保持原始顺序）。"""
    for child in parent.element.body:
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _guess_heading_level(style_name: str) -> Optional[int]:
    match = re.match(r"Heading\s*(\d)", style_name, re.IGNORECASE)
    if match:
        return int(match.group(1))
    match = re.match(r"标题\s*(\d)", style_name)
    if match:
        return int(match.group(1))
    return None


def _get_list_info(para: Paragraph) -> Optional[Dict[str, Any]]:
    """若该段落属于某个 Word 列表，返回 { "num_id": int, "level": int }，否则返回 None。"""
    try:
        p_elm = para._p
        pPr = p_elm.pPr
        if pPr is None:
            return None
        numPr = pPr.numPr
        if numPr is None:
            return None
        num_id = numPr.numId.val if numPr.numId is not None else None
        ilvl = numPr.ilvl.val if numPr.ilvl is not None else 0
        if num_id is not None:
            return {"num_id": num_id, "level": ilvl}
    except Exception:
        pass
    return None


def _detect_manual_list_item(text: str) -> Optional[Dict[str, Any]]:
    """检测手动编号的列表项（非Word列表功能）。
    支持格式: (1), (2), 1., 2., 1), 2), a), b) 等
    返回: {"number": str, "text": str, "format": str} 或 None
    """
    # 匹配模式: 行首的编号
    patterns = [
        (r"^\((\d+)\)\s+(.+)$", "parenthesis"),  # (1) text
        (r"^(\d+)\.\s+(.+)$", "dot"),              # 1. text
        (r"^(\d+)\)\s+(.+)$", "right_paren"),      # 1) text
        (r"^([a-z])\)\s+(.+)$", "letter"),         # a) text
        (r"^([ivxIVX]+)\.\s+(.+)$", "roman"),      # i. text (罗马数字)
    ]
    
    for pattern, fmt in patterns:
        match = re.match(pattern, text, re.DOTALL)
        if match:
            return {
                "number": match.group(1),
                "text": match.group(2).strip(),
                "format": fmt
            }
    return None


def _extract_images_from_paragraph(para: Paragraph) -> List[Dict[str, Any]]:
    """从段落的 runs 中提取图片。使用 a:blip 和 r:embed 精确定位。"""
    images = []
    try:
        for run in para.runs:
            # 使用 local-name() XPath 查询 (python-docx 不支持 namespaces 参数)
            # a:blip 是 DrawingML 中引用图片的地方
            blip_elems = run._r.xpath('.//*[local-name()="blip"]')
            
            # 也要考虑 VML (v:imagedata) - 旧版 Word 兼容
            vml_elems = run._r.xpath('.//*[local-name()="imagedata"]')
            
            embed_ids = []
            for elem in blip_elems:
                # r:embed 属性指向图片的 relationship ID
                rid = elem.get(qn('r:embed'))
                if rid:
                    embed_ids.append(rid)
            for elem in vml_elems:
                rid = elem.get(qn('r:id'))
                if rid:
                    embed_ids.append(rid)
                
            # 去重
            embed_ids = list(set(embed_ids))

            if not embed_ids:
                continue

            # 通过 relationship ID 获取图片 part
            for rid in embed_ids:
                try:
                    part = para.part.related_parts[rid]
                    if hasattr(part, "content_type") and part.content_type.startswith("image/"):
                        path = _save_image(part)
                        if path:
                            images.append({"type": "image", "path": path, "caption": ""})
                except KeyError:
                    continue
    except Exception:
        pass
    return images


def _detect_formula_block(text: str) -> Optional[str]:
    """检测整段是否为块级公式。"""
    text = text.strip()
    # 检测 $$...$$ 或 \[...\]
    if re.match(r"^\s*(\$\$[\s\S]+?\$\$|\\\[[\s\S]+?\\\])\s*$", text):
        # 提取纯 LaTeX
        if text.startswith("$$") and text.endswith("$$"):
            return text[2:-2].strip()
        elif text.startswith("\\[") and text.endswith("\\]"):
            return text[2:-2].strip()
    return None


def _extract_inline_formulas(text: str) -> List[str]:
    """提取行内公式 $...$。"""
    return re.findall(r"\$(?!\$)([^$]+)\$(?!\$)", text)


def _extract_reference_markers(text: str) -> List[str]:
    """提取引用标记，支持数字型和作者-年份型。"""
    markers = []
    # 数字型: [1], [2,3], [1-3] 等，排除任务列表 [ ] [x]
    markers.extend(re.findall(r"\[(?![\sxX\*]\])([0-9,\-\s]+)\]", text))
    # 作者-年份型: (Smith, 2020) (Smith et al., 2020a)
    markers.extend(re.findall(r"\([A-Za-z][^)]*\d{4}[a-z]?\)", text))
    return markers


def _build_inlines_from_runs(para: Paragraph) -> List[Dict[str, Any]]:
    """从段落的 runs 构建行内样式结构。"""
    inlines = []
    for run in para.runs:
        t = run.text
        if not t:
            continue
        inline_item: Dict[str, Any] = {"type": "text", "text": t}
        if run.bold:
            inline_item["bold"] = True
        if run.italic:
            inline_item["italic"] = True
        if run.underline:
            inline_item["underline"] = True
        inlines.append(inline_item)
    return inlines


def _parse_metadata(document: DocxDocument) -> Dict[str, Any]:
    metadata = {"title": "", "authors": [], "abstract": "", "keywords": []}

    # 标题：优先使用 Title 样式；否则第一个非空段落
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style and para.style.name.lower().startswith("title"):
            metadata["title"] = text
            break
        if not metadata["title"]:
            metadata["title"] = text
            break

    author_candidates: List[str] = []
    reached_body = False
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if text.lower().startswith(("abstract", "摘要")):
            reached_body = True
            break
        if text == metadata["title"]:
            continue
        if len(text.split()) <= 20:
            author_candidates.append(text)

    for candidate in author_candidates:
        segments = re.split(r";|、|，|,|\n", candidate)
        for segment in segments:
            segment = segment.strip()
            if len(segment) < 2:
                continue
            name_match = re.match(r"([^(]+)(?:\(([^)]+)\))?", segment)
            if name_match:
                name = name_match.group(1).strip()
                affiliation = name_match.group(2).strip() if name_match.group(2) else ""
                metadata["authors"].append(
                    {"name": name, "affiliation": affiliation, "email": ""}
                )

    abstract_lines: List[str] = []
    keywords: List[str] = []
    collecting_abstract = False
    collecting_keywords = False
    keywords_section_consumed = False
    for para in document.paragraphs:
        text = para.text.strip()
        lower_text = text.lower()
        if not text:
            continue
        if lower_text.startswith("abstract") or lower_text.startswith("摘要"):
            collecting_abstract = True
            collecting_keywords = False
            content = text.split(":", 1)[-1].strip() if ":" in text else ""
            if content:
                abstract_lines.append(content)
            continue
        # 支持更多关键词格式: Keywords, Index Terms, Key words等
        if lower_text.startswith(("keywords", "关键词", "index terms", "key words")):
            # 修复：只处理当前段落，不读取后续内容
            # 尝试多种分隔符: : - —
            keyword_part = ""
            for sep in [":", "-", "—"]:
                if sep in text:
                    keyword_part = text.split(sep, 1)[-1].strip()
                    break
            if not keyword_part:
                keyword_part = text
            # 移除可能的行尾标点
            if keyword_part.endswith('.'):
                keyword_part = keyword_part[:-1]
            keywords.extend([k.strip() for k in re.split(r"[;,，；]", keyword_part) if k.strip()])
            # 关键词提取完成后立即终止循环
            break
        if collecting_abstract:
            abstract_lines.append(text)
            continue
        if collecting_keywords:
            # 这个分支现在不应该执行，因为关键词在同一段落内完成
            if not text or len(text.split()) > 12:
                collecting_keywords = False
                keywords_section_consumed = True
                continue
            keywords.extend([k.strip() for k in re.split(r"[;,，；]", text) if k.strip()])
            continue
        if keywords_section_consumed:
            break

    metadata["abstract"] = "\n".join(abstract_lines).strip()
    metadata["keywords"] = keywords

    if not metadata["title"]:
        raise ValueError("无法从文档中识别标题。")
    if not metadata["authors"]:
        raise ValueError("无法从文档中识别作者信息。")
    if not metadata["abstract"]:
        raise ValueError("无法从文档中识别摘要内容。")

    return metadata


def _paragraph_to_content(para: Paragraph) -> Optional[Dict[str, Any]]:
    text = para.text.strip()
    if not text:
        return None

    # 1. 检测标题
    if para.style:
        heading_level = _guess_heading_level(para.style.name)
        if heading_level:
            # 提取编号和纯文本
            m = re.match(r"^(\d+(\.\d+)*)\s+(.*)$", text)
            if m:
                number = m.group(1)
                pure_text = m.group(3)
            else:
                number = ""
                pure_text = text
            return {
                "type": "heading",
                "level": heading_level,
                "number": number,
                "text": pure_text,
            }
        if "code" in para.style.name.lower():
            return {"type": "code", "language": "", "content": text}

    # 2. 检测块级公式
    formula_block = _detect_formula_block(text)
    if formula_block:
        return {"type": "formula_block", "latex": formula_block}

    # 3. 检测行内公式
    inline_formulas = _extract_inline_formulas(text)
    if inline_formulas:
        # 如果整段几乎都是公式，可以作为多个 formula_inline 返回
        # 这里简化处理：如果有行内公式，记录在段落中
        pass

    # 4. 构建段落，包含行内样式和引用标记
    inlines = _build_inlines_from_runs(para)
    reference_markers = _extract_reference_markers(text)

    result: Dict[str, Any] = {
        "type": "paragraph",
        "text": text,
    }
    if inlines:
        result["inlines"] = inlines
    if reference_markers:
        result["reference_markers"] = reference_markers
    if inline_formulas:
        result["inline_formulas"] = inline_formulas

    return result


def _table_to_content(table: Table) -> Dict[str, Any]:
    table_data: List[List[str]] = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        table_data.append(row_data)
    return {
        "type": "table",
        "caption": "",
        "latex": "",
        "image_path": "",
        "data": table_data,
    }


def _pop_caption_candidate(content_list: List[Dict[str, Any]]) -> str:
    """若上一段落是表格标题则返回并移除，避免重复。"""
    if not content_list:
        return ""
    last_item = content_list[-1]
    if last_item.get("type") != "paragraph":
        return ""
    text = last_item.get("text", "").strip()
    match = CAPTION_PATTERN.match(text)
    if match:
        content_list.pop()
        return text
    return ""


def _find_caption_for_figure_or_table(content_list: List[Dict[str, Any]], search_ahead: int = 3) -> str:
    """查找图表的标题，支持向前和向后搜索。
    
    Args:
        content_list: 当前内容列表
        search_ahead: 向后搜索的段落数量
    
    Returns:
        找到的标题文本，如果没有则返回空字符串
    """
    # 先尝试向前查找（原有逻辑）
    caption = _pop_caption_candidate(content_list)
    if caption:
        return caption
    
    # 如果向前没找到，标记需要向后查找
    # 这个功能需要在主循环中配合实现
    return ""


def parse_docx_to_json(file_path: str) -> Dict[str, Any]:
    document = Document(file_path)
    parsed_data: Dict[str, Any] = {
        "metadata": _parse_metadata(document),
        "content": [],
        "bibliography": [],
    }

    bibliography_started = False
    current_list_items: List[Dict[str, Any]] = []
    current_list_num_id: Optional[int] = None

    def flush_list():
        """将当前累积的列表项输出为一个 list 结构。"""
        nonlocal current_list_items, current_list_num_id
        if current_list_items:
            parsed_data["content"].append(
                {
                    "type": "list",
                    "ordered": True,  # 简化处理，暂时都标记为有序
                    "items": current_list_items,
                }
            )
            current_list_items = []
            current_list_num_id = None

    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            # 先检测图片(图片段落可能文本为空)
            images = _extract_images_from_paragraph(block)
            if images:
                for img in images:
                    caption_text = _pop_caption_candidate(parsed_data["content"])
                    if caption_text:
                        img["caption"] = caption_text
                    parsed_data["content"].append(img)
            
            text = block.text.strip()
            if not text:
                continue

            # 检测参考文献区域
            if text.lower().startswith(("references", "参考文献")):
                flush_list()
                bibliography_started = True
                continue

            if bibliography_started:
                # 扩展参考文献编号格式识别
                bib_id = ""
                match = re.match(r"^\[(\d+)\]|^(\d+)[\.))]|\((\d+)\)", text)
                if match:
                    bib_id = match.group(1) or match.group(2) or match.group(3) or ""
                parsed_data["bibliography"].append(
                    {
                        "id": bib_id,
                        "type": "misc",
                        "authors": [],
                        "title": "",
                        "venue": "",
                        "year": "",
                        "raw": text,
                    }
                )
                continue

            # 检测列表（Word原生列表）
            list_info = _get_list_info(block)
            if list_info:
                num_id = list_info["num_id"]
                level = list_info["level"]
                # 如果列表 ID 变化，先 flush 前一个列表
                if current_list_num_id is not None and current_list_num_id != num_id:
                    flush_list()
                current_list_num_id = num_id
                current_list_items.append({"text": text, "level": level})
                continue
            
            # 检测手动编号列表（非Word列表功能）
            manual_list_item = _detect_manual_list_item(text)
            if manual_list_item:
                # 使用特殊的num_id标识手动列表（负数）
                manual_num_id = -1  # 所有手动列表共享一个ID
                if current_list_num_id is not None and current_list_num_id != manual_num_id:
                    flush_list()
                current_list_num_id = manual_num_id
                current_list_items.append({
                    "text": manual_list_item["text"], 
                    "level": 0,
                    "number": manual_list_item["number"]
                })
                continue
            else:
                # 非列表段落，先 flush 列表
                flush_list()

            # 普通段落处理
            content_entry = _paragraph_to_content(block)
            if content_entry:
                parsed_data["content"].append(content_entry)

        elif isinstance(block, Table):
            # 表格前先 flush 列表
            flush_list()
            table_entry = _table_to_content(block)
            caption_text = _pop_caption_candidate(parsed_data["content"])
            if caption_text:
                table_entry["caption"] = caption_text
            parsed_data["content"].append(table_entry)

    # 最后 flush 可能残留的列表
    flush_list()

    if not parsed_data["bibliography"]:
        parsed_data["bibliography"].append(
            {
                "id": "",
                "type": "misc",
                "authors": [],
                "title": "",
                "venue": "",
                "year": "",
                "raw": "",
            }
        )

    return parsed_data

